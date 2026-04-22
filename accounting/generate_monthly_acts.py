"""
Генератор ежемесячных актов выполненных работ (СММ и Услуги).

Вход: Excel-файл на Яндекс.Диске (папка статистики), листы:
  - «СММ посты»  : Клиент | Договор | ДС | Дата | Тип контента | Ссылка | Стоимость
  - «Услуги»     : Клиент | Договор | ДС | ТЗ | Услуга | Стоимость

Выход: DOCX + PDF на Яндекс.Диске в папке Бухгалтерия/{Клиент}/{YYYY-MM}/
"""

import io
import json
import os
import subprocess
import sys
import tempfile
from datetime import date, datetime
from pathlib import Path

import openpyxl
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from num2words import num2words

YADISK_API = 'https://cloud-api.yandex.net/v1/disk'

MONTHS_RU = {
    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
    5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
    9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря',
}

MONTH_NAMES_RU = {
    1: 'январе', 2: 'феврале', 3: 'марте', 4: 'апреле',
    5: 'мае', 6: 'июне', 7: 'июле', 8: 'августе',
    9: 'сентябре', 10: 'октябре', 11: 'ноябре', 12: 'декабре',
}


# ---------------------------------------------------------------------------
# Яндекс.Диск helpers
# ---------------------------------------------------------------------------

def _ya_headers(token):
    return {'Authorization': f'OAuth {token}'}


def ensure_yadisk_folder(token, path):
    parts = [p for p in path.strip('/').split('/') if p]
    cur = ''
    for part in parts:
        cur = f'{cur}/{part}' if cur else part
        r = requests.put(
            f'{YADISK_API}/resources',
            params={'path': cur},
            headers=_ya_headers(token),
            timeout=30,
        )
        if r.status_code not in (201, 409):
            raise RuntimeError(f'mkdir {cur}: {r.status_code} {r.text[:200]}')


def upload_file_to_yadisk(token, local_path, remote_path):
    r = requests.get(
        f'{YADISK_API}/resources/upload',
        params={'path': remote_path, 'overwrite': 'true'},
        headers=_ya_headers(token),
        timeout=30,
    )
    r.raise_for_status()
    href = r.json()['href']
    with open(local_path, 'rb') as fh:
        put = requests.put(href, data=fh, timeout=300)
    if put.status_code >= 400:
        raise RuntimeError(f'PUT {put.status_code}: {put.text[:200]}')
    print(f'  ✓ загружен: {remote_path}')


def download_excel_from_yadisk(token, remote_path):
    """Скачивает Excel из приватной папки, возвращает bytes."""
    r = requests.get(
        f'{YADISK_API}/resources/download',
        params={'path': remote_path},
        headers=_ya_headers(token),
        timeout=30,
    )
    if r.status_code == 404:
        raise FileNotFoundError(f'Файл не найден на Яндекс.Диске: {remote_path}')
    r.raise_for_status()
    data = requests.get(r.json()['href'], timeout=120)
    data.raise_for_status()
    return data.content


def list_yadisk_folder(token, path):
    """Возвращает список файлов в папке."""
    r = requests.get(
        f'{YADISK_API}/resources',
        params={'path': path, 'limit': 100, 'fields': '_embedded.items.name,_embedded.items.type'},
        headers=_ya_headers(token),
        timeout=30,
    )
    if r.status_code == 404:
        return []
    r.raise_for_status()
    return r.json().get('_embedded', {}).get('items', [])


# ---------------------------------------------------------------------------
# Config helpers
# ---------------------------------------------------------------------------

def load_config():
    cfg_path = Path(__file__).parent / 'config' / 'clients.json'
    with open(cfg_path, encoding='utf-8') as f:
        return json.load(f)


def find_client_by_name(cfg, name):
    """Ищет клиента по display_name или folder_name (частичное совпадение)."""
    name_lower = name.strip().lower()
    for c in cfg['clients']:
        if (name_lower in c['display_name'].lower() or
                name_lower in c['folder_name'].lower() or
                name_lower in c['id'].lower()):
            return c
    return None


def get_ds_info(client, ds_number):
    """Возвращает дополнительное соглашение по номеру."""
    for contract in client['contracts']:
        for ds in contract['additional_agreements']:
            if str(ds['number']) == str(ds_number):
                return contract, ds
    return None, None


# ---------------------------------------------------------------------------
# Сумма прописью
# ---------------------------------------------------------------------------

def amount_in_words(amount):
    rubles = int(amount)
    kopecks = round((amount - rubles) * 100)
    words = num2words(rubles, lang='ru', to='currency',
                      currency='RUB', separator=' ').strip()
    words = words[0].upper() + words[1:]
    if kopecks == 0:
        return f'{words} 00 копеек'
    return f'{words} {kopecks:02d} копеек'


# ---------------------------------------------------------------------------
# DOCX — форматирование
# ---------------------------------------------------------------------------

def _set_font(run, size=11, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Times New Roman'


def _para(doc, text='', align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p


def _set_margins(doc, top=2, bottom=2, left=3, right=1.5):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


# ---------------------------------------------------------------------------
# Генерация АКТ СММ (тип «smm»)
# ---------------------------------------------------------------------------

def build_smm_act(cfg, client, contract, ds, rows, period_year, period_month,
                  act_number, ds_tz_number=None):
    """
    rows: список dict {date_str, content_type, link, amount}
    Возвращает путь к временному DOCX-файлу.
    """
    contractor = cfg['contractor']
    smm_rates = cfg['smm_rates']
    last_day = _last_day_of_month(period_year, period_month)
    month_gen = MONTHS_RU[period_month]
    month_in = MONTH_NAMES_RU[period_month]

    # Считаем стоимость автоматически если не задана
    for row in rows:
        if not row.get('amount'):
            row['amount'] = smm_rates.get(row['content_type'], 0)

    total = sum(r['amount'] for r in rows)

    doc = Document()
    _set_margins(doc)

    # Заголовок
    _para(doc, 'АКТ ПРИЕМКИ ВЫПОЛНЕННЫХ РАБОТ',
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    _para(doc)

    ds_title = (
        f'К Дополнительному соглашению № {ds["number"]} от {ds["date"]} г. '
        f'к Договору № {contract["number"]} {contract["title"]} от {contract["date"]} г.'
    )
    _para(doc, ds_title, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _para(doc)

    # Дата и место
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run('г. Пермь')
    _set_font(r1)
    tab = p.add_run('\t\t\t\t\t')
    _set_font(tab)
    r2 = p.add_run(f'«{last_day}» {month_gen} {period_year} г.')
    _set_font(r2)
    _para(doc)

    # Преамбула
    preambula = (
        f'Заказчик {client["full_name"]}, именуемый в дальнейшем «Заказчик», '
        f'с одной стороны, и Подрядчик {contractor["name"]}, именуемая в дальнейшем '
        f'«Подрядчик», с другой стороны, совместно именуемые «Стороны», составили '
        f'настоящий Акт о следующем:'
    )
    _para(doc, preambula, size=11)
    _para(doc)
    _para(doc,
          f'Для оказания услуг по разработке и продвижению бренда Заказчика, '
          f'Стороны {ds["date"]} года заключили Дополнительное соглашение '
          f'№{ds["number"]} к Договору № {contract["number"]} от {contract["date"]} г.',
          size=11)
    _para(doc)

    _para(doc, '3. Перечень выполненных работ', bold=True, size=11)
    _para(doc,
          f'В отчетном периоде с 1 по {last_day} {month_gen} {period_year} г. '
          f'Подрядчиком выполнены и опубликованы следующие единицы контента '
          f'в социальных сетях Заказчика:',
          size=11)
    _para(doc)

    # Таблица публикаций
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, title in enumerate(['№', 'Дата публикации', 'Тип контента', 'Стоимость, ₽', 'Ссылка']):
        hdr[i].text = title
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)

    for idx, row in enumerate(rows, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = row.get('date_str', '')
        cells[2].text = row.get('content_type', '')
        cells[3].text = str(int(row['amount']))
        cells[4].text = row.get('link', '')
        for cell in cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    _para(doc)
    _para(doc,
          f'Общая стоимость фактически оказанных услуг за отчетный период составляет: '
          f'{int(total):,} ({amount_in_words(total)}) рублей 00 копеек, НДС не облагается.'.replace(',', ' '),
          size=11)
    _para(doc)

    closing = (
        'На основании изложенного, Стороны заявляют, что Работы выполнены согласно '
        'условиям Договора в полном объеме, надлежащего качества, претензий друг к другу '
        'по исполнению Договора Стороны не имеют.\n\n'
        f'В соответствии с условиями Договора оплата Подрядчику за Работы по '
        f'Дополнительному соглашению №{ds["number"]} от {ds["date"]} г. к '
        f'Договору №{contract["number"]} от {contract["date"]} г. произведена '
        f'Заказчиком в полном объёме.\n\n'
        'Настоящий акт выполнения работ составлен в двух экземплярах, имеющих '
        'одинаковую юридическую силу, по одному экземпляру для каждой из Сторон.'
    )
    _para(doc, closing, size=11)
    _para(doc)

    # Реквизиты
    _add_signatures(doc, contractor, client)

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    return tmp.name


# ---------------------------------------------------------------------------
# Генерация АКТ по Услугам (тип «services») + ТЗ
# ---------------------------------------------------------------------------

def build_services_tz(cfg, client, contract, ds, services, period_year, period_month, tz_number):
    """ТЗ (Техническое задание). Возвращает путь к DOCX."""
    contractor = cfg['contractor']
    last_day = _last_day_of_month(period_year, period_month)
    month_gen = MONTHS_RU[period_month]
    total = sum(s['amount'] for s in services)

    doc = Document()
    _set_margins(doc)

    _para(doc, f'ТЕХНИЧЕСКОЕ ЗАДАНИЕ № {tz_number}',
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    _para(doc,
          f'Приложение к Дополнительному соглашению № {ds["number"]} '
          f'от {ds["date"]} г. к Договору № {contract["number"]} '
          f'от {contract["date"]} г.',
          align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _para(doc)
    _para(doc, f'Отчётный период: {month_gen} {period_year} г.', size=11)
    _para(doc, f'г. Пермь   «{last_day}» {month_gen} {period_year} г.', size=11)
    _para(doc)
    _para(doc, 'Перечень услуг:', bold=True, size=11)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, t in enumerate(['№', 'Наименование услуги', 'Стоимость, ₽']):
        table.rows[0].cells[i].text = t
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for idx, s in enumerate(services, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = s['name']
        cells[2].text = f'{int(s["amount"]):,}'.replace(',', ' ')

    # Итого
    cells = table.add_row().cells
    cells[1].text = 'ИТОГО:'
    cells[1].paragraphs[0].runs[0].font.bold = True
    cells[2].text = f'{int(total):,}'.replace(',', ' ')
    cells[2].paragraphs[0].runs[0].font.bold = True

    _para(doc)
    _add_signatures(doc, contractor, client)

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    return tmp.name


def build_services_act(cfg, client, contract, ds, services, period_year, period_month,
                       act_number, tz_number):
    contractor = cfg['contractor']
    last_day = _last_day_of_month(period_year, period_month)
    month_gen = MONTHS_RU[period_month]
    total = sum(s['amount'] for s in services)

    doc = Document()
    _set_margins(doc)

    _para(doc, 'АКТ ПРИЕМКИ ВЫПОЛНЕННЫХ РАБОТ',
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    _para(doc,
          f'К Дополнительному соглашению № {ds["number"]} от {ds["date"]} г. '
          f'к Договору № {contract["number"]} {contract["title"]} от {contract["date"]} г.',
          align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _para(doc)

    p = doc.add_paragraph()
    p.add_run('г. Пермь').font.size = Pt(11)
    p.add_run(f'\t\t\t\t\t«{last_day}» {month_gen} {period_year} г.').font.size = Pt(11)
    _para(doc)

    preambula = (
        f'Заказчик {client["full_name"]}, именуемый в дальнейшем «Заказчик», '
        f'с одной стороны, и Подрядчик {contractor["name"]}, именуемая в дальнейшем '
        f'«Подрядчик», с другой стороны, составили настоящий Акт о следующем:\n\n'
        f'Для оказания услуг Стороны заключили Дополнительное соглашение №{ds["number"]} '
        f'к Договору №{contract["number"]} от {contract["date"]} г. и Приложение к нему – '
        f'Техническое задание №{tz_number} (далее – Техническое задание).\n\n'
        f'Подрядчик провел Работы в соответствии с вышеуказанным Техническим заданием №{tz_number}, а именно:'
    )
    _para(doc, preambula, size=11)
    _para(doc)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, t in enumerate(['№', 'Наименование', 'Стоимость']):
        table.rows[0].cells[i].text = t
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for idx, s in enumerate(services, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = s['name']
        cells[2].text = f'{int(s["amount"]):,} рублей'.replace(',', ' ')

    cells = table.add_row().cells
    cells[1].text = 'ИТОГО:'
    cells[1].paragraphs[0].runs[0].font.bold = True
    cells[2].text = f'{int(total):,} рублей'.replace(',', ' ')
    cells[2].paragraphs[0].runs[0].font.bold = True

    _para(doc)
    _para(doc,
          f'На основании изложенного, Стороны заявляют, что Работы по указанному выше '
          f'Техническому заданию выполнены согласно условиям Договора в полном объеме, '
          f'надлежащего качества, претензий друг к другу по исполнению Договора Стороны не имеют.\n\n'
          f'В соответствии с условиями Договора оплата Подрядчику за Работы по '
          f'Дополнительному соглашению №{ds["number"]} от {ds["date"]} г. к '
          f'Договору №{contract["number"]} от {contract["date"]} г. произведена '
          f'Заказчиком в полном объёме.\n\n'
          f'Настоящий акт выполнения работ составлен в двух экземплярах, имеющих '
          f'одинаковую юридическую силу, по одному экземпляру для каждой из Сторон.',
          size=11)
    _para(doc)
    _add_signatures(doc, contractor, client)

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    return tmp.name


# ---------------------------------------------------------------------------
# Блок реквизитов сторон
# ---------------------------------------------------------------------------

def _add_signatures(doc, contractor, client):
    _para(doc, 'Реквизиты и подписи сторон.', bold=True, size=11)
    _para(doc)

    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells

    def _block(lines):
        return '\n'.join(lines)

    left.text = _block([
        'Подрядчик:',
        contractor['name'],
        f'ИНН: {contractor["inn"]}',
        f'ОГРНИП: {contractor["ogrnip"]}',
        f'Юр. адрес: {contractor["address"]}',
        f'Р/c: {contractor["rs"]}',
        f'К/c: {contractor["ks"]}',
        f'Банк {contractor["bank"]}',
        f'БИК банка {contractor["bik"]}',
        f'Email: {contractor["email"]}',
        f'Тел.: {contractor["phone"]}',
        '',
        '_______________/ИП Ладыгина В.А./',
    ])

    right.text = _block([
        'Заказчик:',
        client['full_name'],
        f'ИНН {client["inn"]}',
        f'ОГРНИП {client["ogrnip"]},',
        f'Юр.адрес: {client["address"]}.',
        f'р/с {client["rs"]} в {client["bank"]}',
        f'к/с {client["ks"]}',
        f'БИК {client["bik"]}',
        '',
        f'___________________/{client["short_name"].split()[-1][0]}. {client["short_name"].split()[1][0]}. {client["short_name"].split()[0]}/',
    ])

    for cell in [left, right]:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# DOCX → PDF через LibreOffice
# ---------------------------------------------------------------------------

def docx_to_pdf(docx_path, output_dir=None):
    if output_dir is None:
        output_dir = str(Path(docx_path).parent)
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f'LibreOffice PDF конвертация не удалась:\n{result.stderr}')
    pdf_path = str(Path(docx_path).with_suffix('.pdf'))
    if not Path(pdf_path).exists():
        # LibreOffice иногда сохраняет в output_dir с тем же именем
        pdf_path = str(Path(output_dir) / Path(docx_path).with_suffix('.pdf').name)
    return pdf_path


# ---------------------------------------------------------------------------
# Чтение входного Excel
# ---------------------------------------------------------------------------

def parse_smm_sheet(wb):
    """Читает лист «СММ посты», возвращает список dict."""
    if 'СММ посты' not in wb.sheetnames:
        return []
    ws = wb['СММ посты']
    rows = []
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 2):
        if not row[0]:
            continue
        client_name, contract_num, ds_num, date_val, content_type, link, amount = (
            row[0], row[1], row[2], row[3], row[4], row[5],
            row[6] if len(row) > 6 else None
        )
        if isinstance(date_val, datetime):
            date_str = date_val.strftime('%d.%m.%Yг.')
        else:
            date_str = str(date_val) if date_val else ''
        rows.append({
            'client_name': str(client_name).strip(),
            'contract_num': str(contract_num).strip(),
            'ds_num': str(ds_num).strip(),
            'date_str': date_str,
            'content_type': str(content_type).strip() if content_type else '',
            'link': str(link).strip() if link else '',
            'amount': float(amount) if amount else None,
        })
    return rows


def parse_services_sheet(wb):
    """Читает лист «Услуги», возвращает список dict."""
    if 'Услуги' not in wb.sheetnames:
        return []
    ws = wb['Услуги']
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row[0]:
            continue
        client_name, contract_num, ds_num, tz_num, service_name, amount = (
            row[0], row[1], row[2], row[3], row[4], row[5]
        )
        rows.append({
            'client_name': str(client_name).strip(),
            'contract_num': str(contract_num).strip(),
            'ds_num': str(ds_num).strip(),
            'tz_num': str(tz_num).strip(),
            'name': str(service_name).strip(),
            'amount': float(amount) if amount else 0.0,
        })
    return rows


# ---------------------------------------------------------------------------
# Вспомогательные
# ---------------------------------------------------------------------------

def _last_day_of_month(year, month):
    if month == 12:
        return 31
    return (date(year, month + 1, 1) - date(year, month, 1)).days + \
           date(year, month, 1).day - 1


def _last_day_of_month(year, month):
    import calendar
    return calendar.monthrange(year, month)[1]


# ---------------------------------------------------------------------------
# Основная функция
# ---------------------------------------------------------------------------

def run(period_year, period_month, token, accounting_folder, stats_excel_path, dry_run=False):
    cfg = load_config()
    period_str = f'{period_year}-{period_month:02d}'
    print(f'\n=== Генерация актов за {period_str} ===')

    # Скачиваем Excel
    print(f'Скачиваем данные: {stats_excel_path}')
    excel_bytes = download_excel_from_yadisk(token, stats_excel_path)
    wb = openpyxl.load_workbook(io.BytesIO(excel_bytes))

    smm_rows = parse_smm_sheet(wb)
    services_rows = parse_services_sheet(wb)
    print(f'  СММ строк: {len(smm_rows)},  Услуги строк: {len(services_rows)}')

    generated = []

    # --- СММ акты ---
    smm_by_client_ds = {}
    for row in smm_rows:
        key = (row['client_name'], row['ds_num'])
        smm_by_client_ds.setdefault(key, []).append(row)

    for (client_name, ds_num), rows in smm_by_client_ds.items():
        client = find_client_by_name(cfg, client_name)
        if not client:
            print(f'  ⚠ Клиент не найден: "{client_name}" — пропускаем')
            continue
        contract, ds = get_ds_info(client, ds_num)
        if not ds:
            print(f'  ⚠ ДС №{ds_num} не найдено для {client_name} — пропускаем')
            continue

        act_num = cfg['act_counters']['last_act_number'] + 1
        total = sum(r['amount'] or cfg['smm_rates'].get(r['content_type'], 0) for r in rows)
        print(f'\n  Клиент: {client["display_name"]}  ДС №{ds_num}')
        print(f'  Постов: {len(rows)},  Сумма: {int(total):,} руб.'.replace(',', ' '))

        if dry_run:
            print(f'  [dry-run] АКТ №{act_num} будет создан')
            continue

        docx_path = build_smm_act(cfg, client, contract, ds, rows,
                                  period_year, period_month, act_num)
        pdf_path = docx_to_pdf(docx_path)

        remote_dir = f'{accounting_folder}/{client["folder_name"]}/{period_str}'
        ensure_yadisk_folder(token, remote_dir)

        prefix = ds.get('act_prefix', f'АКТ_СММ_ДС{ds_num}')
        upload_file_to_yadisk(token, docx_path, f'{remote_dir}/{prefix}_{period_str}.docx')
        upload_file_to_yadisk(token, pdf_path, f'{remote_dir}/{prefix}_{period_str}.pdf')

        generated.append({
            'client': client['display_name'], 'type': 'СММ', 'ds': ds_num,
            'total': total, 'act_num': act_num,
            'path': f'{remote_dir}/{prefix}_{period_str}.pdf',
        })
        Path(docx_path).unlink(missing_ok=True)
        Path(pdf_path).unlink(missing_ok=True)

    # --- Акты по Услугам ---
    svc_by_client_ds_tz = {}
    for row in services_rows:
        key = (row['client_name'], row['ds_num'], row['tz_num'])
        svc_by_client_ds_tz.setdefault(key, []).append(row)

    for (client_name, ds_num, tz_num), svc_rows in svc_by_client_ds_tz.items():
        client = find_client_by_name(cfg, client_name)
        if not client:
            print(f'  ⚠ Клиент не найден: "{client_name}" — пропускаем')
            continue
        contract, ds = get_ds_info(client, ds_num)
        if not ds:
            print(f'  ⚠ ДС №{ds_num} не найдено для {client_name} — пропускаем')
            continue

        act_num = cfg['act_counters']['last_act_number'] + 1
        services = [{'name': r['name'], 'amount': r['amount']} for r in svc_rows]
        total = sum(s['amount'] for s in services)
        print(f'\n  Клиент: {client["display_name"]}  ДС №{ds_num}  ТЗ №{tz_num}')
        print(f'  Услуг: {len(services)},  Сумма: {int(total):,} руб.'.replace(',', ' '))

        if dry_run:
            print(f'  [dry-run] ТЗ №{tz_num} + АКТ №{act_num} будут созданы')
            continue

        tz_path = build_services_tz(cfg, client, contract, ds, services,
                                    period_year, period_month, tz_num)
        act_path = build_services_act(cfg, client, contract, ds, services,
                                      period_year, period_month, act_num, tz_num)
        tz_pdf = docx_to_pdf(tz_path)
        act_pdf = docx_to_pdf(act_path)

        remote_dir = f'{accounting_folder}/{client["folder_name"]}/{period_str}'
        ensure_yadisk_folder(token, remote_dir)

        prefix = ds.get('act_prefix', f'АКТ_ДС{ds_num}')
        upload_file_to_yadisk(token, tz_path,  f'{remote_dir}/ТЗ_{tz_num}_{period_str}.docx')
        upload_file_to_yadisk(token, tz_pdf,   f'{remote_dir}/ТЗ_{tz_num}_{period_str}.pdf')
        upload_file_to_yadisk(token, act_path, f'{remote_dir}/{prefix}_{period_str}.docx')
        upload_file_to_yadisk(token, act_pdf,  f'{remote_dir}/{prefix}_{period_str}.pdf')

        generated.append({
            'client': client['display_name'], 'type': 'Услуги', 'ds': ds_num,
            'tz': tz_num, 'total': total, 'act_num': act_num,
            'path': f'{remote_dir}/{prefix}_{period_str}.pdf',
        })
        for p in [tz_path, act_path, tz_pdf, act_pdf]:
            Path(p).unlink(missing_ok=True)

    return generated


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--month', required=True, help='YYYY-MM')
    parser.add_argument('--stats-path', required=True,
                        help='Путь на Яндекс.Диске к Excel-файлу статистики')
    parser.add_argument('--dry-run', action='store_true')
    args = parser.parse_args()

    year, month = map(int, args.month.split('-'))
    token = os.environ['YANDEX_DISK_TOKEN']
    folder = os.environ.get('YANDEX_DISK_ACCOUNTING_FOLDER', 'Бухгалтерия')

    docs = run(year, month, token, folder, args.stats_path, dry_run=args.dry_run)
    print(f'\nГотово: {len(docs)} документ(ов)')
    for d in docs:
        print(f'  {d["client"]} | ДС №{d["ds"]} | {d["type"]} | '
              f'{int(d["total"]):,} руб. → {d["path"]}'.replace(',', ' '))
