"""
Генератор шаблона статистики для ВК ОРД и инструкции по заполнению.

Вход: лист «ОРД статистика» входного Excel (или отдельный файл).
  Колонки: ERID | Название | Площадка | Показов | Оплачено показов |
           Период с | Период по | Тип события | Стоимость события | Сумма с НДС | Ставка НДС

Выход:
  1. Шаблон_ОРД_{YYYY-MM}.xlsx  — заполненный шаблон для сверки/ввода в ВК ОРД
  2. Инструкция_ОРД.docx/.pdf   — пошаговая инструкция
"""

import io
import json
import os
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

YADISK_API = 'https://cloud-api.yandex.net/v1/disk'

# Поля формы ВК ОРД → маппинг из входного Excel
ORD_COLUMNS = [
    'ERID (токен)',
    'Название креатива',
    'Площадка',
    'Количество показов',
    'Оплаченное кол-во показов',
    'Период с (дата начала)',
    'Период по (дата конца)',
    'Тип платного события',
    'Стоимость одного события, ₽',
    'Сумма (с НДС, если облагается), ₽',
    'Ставка НДС, %',
]

VK_ORD_FIELDS_DESCRIPTION = {
    'ERID (токен)': 'Уникальный идентификатор рекламного материала (erid). Выдаётся при регистрации креатива в ВК ОРД.',
    'Название креатива': 'Название из карточки креатива в ВК ОРД (для сверки).',
    'Площадка': 'Название площадки из раздела «Площадки» в ВК ОРД.',
    'Количество показов': 'Фактическое количество показов рекламного материала за период.',
    'Оплаченное кол-во показов': 'Количество показов, за которые произведена оплата (обычно = фактическим).',
    'Период с (дата начала)': 'Фактическая дата начала показа креатива (ДД.ММ.ГГГГ).',
    'Период по (дата конца)': 'Фактическая дата окончания показа (ДД.ММ.ГГГГ).',
    'Тип платного события': 'Выбирается из списка: CPM (за 1000 показов), CPC (за клик), CPA (за действие), Фиксированная.',
    'Стоимость одного события, ₽': 'Цена одного события (показ / клик / действие).',
    'Сумма (с НДС, если облагается), ₽': 'Итоговая сумма по данному размещению.',
    'Ставка НДС, %': 'Ставка НДС: 20, 10 или 0. Если не облагается — оставить пустым.',
}


# ---------------------------------------------------------------------------
# Яндекс.Диск helpers (дублируются для автономности модуля)
# ---------------------------------------------------------------------------

def _ya_headers(token):
    return {'Authorization': f'OAuth {token}'}


def ensure_yadisk_folder(token, path):
    parts = [p for p in path.strip('/').split('/') if p]
    cur = ''
    for part in parts:
        cur = f'{cur}/{part}' if cur else part
        r = requests.put(
            f'{YADISK_API}/resources',
            params={'path': cur},
            headers=_ya_headers(token),
            timeout=30,
        )
        if r.status_code not in (201, 409):
            raise RuntimeError(f'mkdir {cur}: {r.status_code} {r.text[:200]}')


def upload_file_to_yadisk(token, local_path, remote_path):
    r = requests.get(
        f'{YADISK_API}/resources/upload',
        params={'path': remote_path, 'overwrite': 'true'},
        headers=_ya_headers(token),
        timeout=30,
    )
    r.raise_for_status()
    href = r.json()['href']
    with open(local_path, 'rb') as fh:
        put = requests.put(href, data=fh, timeout=300)
    if put.status_code >= 400:
        raise RuntimeError(f'PUT {put.status_code}: {put.text[:200]}')
    print(f'  ✓ загружен: {remote_path}')


def download_excel_from_yadisk(token, remote_path):
    r = requests.get(
        f'{YADISK_API}/resources/download',
        params={'path': remote_path},
        headers=_ya_headers(token),
        timeout=30,
    )
    if r.status_code == 404:
        raise FileNotFoundError(f'Файл не найден: {remote_path}')
    r.raise_for_status()
    data = requests.get(r.json()['href'], timeout=120)
    data.raise_for_status()
    return data.content


# ---------------------------------------------------------------------------
# Чтение листа «ОРД статистика»
# ---------------------------------------------------------------------------

def parse_ord_sheet(wb):
    sheet_name = 'ОРД статистика'
    if sheet_name not in wb.sheetnames:
        print(f'  Лист «{sheet_name}» не найден в Excel — пропускаем ОРД')
        return []
    ws = wb[sheet_name]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row[0]:
            continue
        erid, name, platform, shows, paid_shows, date_from, date_to, \
            event_type, event_cost, total_sum, vat = (
                list(row) + [None] * 11
            )[:11]

        def fmt_date(v):
            if isinstance(v, datetime):
                return v.strftime('%d.%m.%Y')
            return str(v) if v else ''

        rows.append({
            'erid': str(erid).strip() if erid else '',
            'name': str(name).strip() if name else '',
            'platform': str(platform).strip() if platform else '',
            'shows': int(shows) if shows else 0,
            'paid_shows': int(paid_shows) if paid_shows else (int(shows) if shows else 0),
            'date_from': fmt_date(date_from),
            'date_to': fmt_date(date_to),
            'event_type': str(event_type).strip() if event_type else '',
            'event_cost': float(event_cost) if event_cost else None,
            'total_sum': float(total_sum) if total_sum else None,
            'vat': str(vat).strip() if vat else '',
        })
    return rows


# ---------------------------------------------------------------------------
# Создание XLSX-шаблона для ОРД
# ---------------------------------------------------------------------------

def build_ord_xlsx(rows, period_str):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Статистика ОРД'

    # Стили
    header_fill = PatternFill('solid', fgColor='1C4587')
    header_font = Font(color='FFFFFF', bold=True, size=10)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Заголовок
    ws.merge_cells('A1:K1')
    ws['A1'] = f'Статистика показов для ВК ОРД — {period_str}'
    ws['A1'].font = Font(bold=True, size=12)
    ws['A1'].alignment = center

    # Строка с инструкцией
    ws.merge_cells('A2:K2')
    ws['A2'] = ('Заполните таблицу и перенесите данные в форму «Добавить статистику» '
                'в разделе ВК ОРД → Статистика. '
                'Поля, помеченные (*), обязательны.')
    ws['A2'].font = Font(italic=True, size=9, color='666666')
    ws['A2'].alignment = Alignment(wrap_text=True)
    ws.row_dimensions[2].height = 30

    # Заголовки колонок
    headers = ORD_COLUMNS
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    ws.row_dimensions[3].height = 40

    # Данные
    for row_idx, row in enumerate(rows, 4):
        values = [
            row['erid'],
            row['name'],
            row['platform'],
            row['shows'],
            row['paid_shows'],
            row['date_from'],
            row['date_to'],
            row['event_type'],
            row['event_cost'],
            row['total_sum'],
            row['vat'],
        ]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.border = border
            cell.alignment = Alignment(vertical='center')

    # Ширина колонок
    widths = [20, 30, 25, 15, 18, 16, 16, 20, 22, 22, 12]
    for col_idx, width in enumerate(widths, 1):
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width

    # Лист «Пояснения»
    ws2 = wb.create_sheet('Пояснения по полям')
    ws2['A1'] = 'Поле'
    ws2['B1'] = 'Описание'
    for cell in [ws2['A1'], ws2['B1']]:
        cell.font = Font(bold=True)
    for r_idx, (field, desc) in enumerate(VK_ORD_FIELDS_DESCRIPTION.items(), 2):
        ws2.cell(row=r_idx, column=1, value=field)
        ws2.cell(row=r_idx, column=2, value=desc)
        ws2.cell(row=r_idx, column=2).alignment = Alignment(wrap_text=True)
    ws2.column_dimensions['A'].width = 30
    ws2.column_dimensions['B'].width = 70

    tmp = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
    wb.save(tmp.name)
    return tmp.name


# ---------------------------------------------------------------------------
# Инструкция по заполнению ВК ОРД (DOCX)
# ---------------------------------------------------------------------------

def build_ord_instruction():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1C, 0x45, 0x87)

    def para(text):
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(11)

    def bullet(text):
        p = doc.add_paragraph(text, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    h('Инструкция по заполнению статистики в ВК ОРД', level=1)
    para('Документ описывает порядок передачи статистики показов '
         'по промаркированным рекламным материалам через кабинет ВК ОРД.')

    h('1. Что такое ВК ОРД и зачем передавать статистику', level=2)
    para('ВК ОРД (Оператор рекламных данных ВКонтакте) — система учёта интернет-рекламы, '
         'обязательная по закону о маркировке рекламы (ФЗ №347). Каждый рекламный материал '
         'должен получить ERID-токен, а по итогам месяца в ОРД передаются данные '
         'о количестве показов и стоимости размещения.')

    h('2. Что нужно подготовить', level=2)
    for item in [
        'Доступ в кабинет ord.vk.com (ВК ID)',
        'ERID-токены для каждого креатива (раздел «Креативы»)',
        'Статистику показов от площадок (сообщества ВК, сайты, OK и т.д.)',
        'Финансовые данные: сумму по каждому размещению',
    ]:
        bullet(item)

    h('3. Пошаговый порядок заполнения формы', level=2)

    steps = [
        ('Перейти в раздел «Статистика»',
         'На главной странице ord.vk.com выберите вкладку «Статистика».'),
        ('Нажать «Добавить статистику»',
         'Кнопка «+ Добавить статистику» в верхнем левом углу.'),
        ('Выбрать Креатив',
         'В выпадающем списке «Креатив» найдите нужный по названию или ERID. '
         'ERID отображается в столбце «ERID» в разделе «Креативы».'),
        ('Выбрать Площадку',
         'В поле «Площадка» выберите площадку, где показывалась реклама. '
         'Площадки добавляются заранее в разделе «Площадки».'),
        ('Указать количество показов',
         'Поле «Количество показов» — суммарное число показов за период. '
         'Поле «Оплаченное количество показов» — оставьте равным фактическому, '
         'если не знаете точное платное количество.'),
        ('Проверить галочку «Плановые даты совпадают с фактическими»',
         'Если флажок установлен — даты плана = даты факта. '
         'Введите «Фактическую дату начала» и «Фактическую дату конца» показа.'),
        ('Заполнить Финансовую информацию',
         'Тип платного события: выберите из списка (чаще всего — «Фиксированная»).\n'
         'Сумма (с НДС): итоговая стоимость размещения.\n'
         'Ставка НДС: выберите 20%, 10% или «Без НДС». ИП на УСН — «Без НДС».\n'
         'Галочка «Рассчитать автоматически» — система сама посчитает НДС из суммы.'),
        ('Сохранить',
         'Нажмите «Сохранить». Запись появится в таблице статистики со статусом '
         '«Ожидает обработки», затем «Принято в ЕРИР».'),
    ]

    for num, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'Шаг {num}: {title}')
        r.bold = True
        r.font.size = Pt(11)
        para(desc)
        doc.add_paragraph()

    h('4. Сроки передачи статистики', level=2)
    para('Статистику за предыдущий месяц необходимо передать не позднее '
         '20-го числа следующего месяца. Рекомендуется делать это с 1 по 3 число '
         'одновременно с формированием актов выполненных работ.')

    h('5. Использование шаблона Excel', level=2)
    para('Файл «Шаблон_ОРД_{месяц}.xlsx» содержит заполненные данные по всем '
         'размещениям. Используйте его для сверки перед вводом в форму ВК ОРД. '
         'Лист «Пояснения по полям» содержит расшифровку каждого поля.')

    h('6. Ошибки и их решение', level=2)
    errors = [
        ('Ошибка «Площадка не найдена»',
         'Добавьте площадку в раздел «Площадки» ВК ОРД перед заполнением статистики.'),
        ('Статус «Отклонено»',
         'Проверьте правильность ERID и соответствие площадки договору.'),
        ('Поле «Сумма» подсвечено красным',
         'Убедитесь, что введена сумма и выбрана ставка НДС.'),
    ]
    for err, solution in errors:
        p = doc.add_paragraph()
        r = p.add_run(f'• {err}: ')
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(solution).font.size = Pt(11)

    h('7. Контакты поддержки ВК ОРД', level=2)
    para('Email: support@ord.vk.com\n'
         'Официальная документация: ord.vk.com/help/\n'
         'API-документация: ord.vk.com/help/api/')

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(tmp.name)
    return tmp.name


# ---------------------------------------------------------------------------
# PDF конвертация
# ---------------------------------------------------------------------------

def docx_to_pdf(docx_path):
    output_dir = str(Path(docx_path).parent)
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f'LibreOffice error:\n{result.stderr}')
    pdf_path = str(Path(docx_path).with_suffix('.pdf'))
    if not Path(pdf_path).exists():
        pdf_path = str(Path(output_dir) / Path(docx_path).with_suffix('.pdf').name)
    return pdf_path


# ---------------------------------------------------------------------------
# Основная функция
# ---------------------------------------------------------------------------

def run(period_year, period_month, token, accounting_folder,
        stats_excel_path, dry_run=False):
    import openpyxl
    period_str = f'{period_year}-{period_month:02d}'
    print(f'\n=== ОРД шаблон за {period_str} ===')

    excel_bytes = download_excel_from_yadisk(token, stats_excel_path)
    wb = openpyxl.load_workbook(io.BytesIO(excel_bytes))
    rows = parse_ord_sheet(wb)

    if not rows:
        print('  Нет данных для ОРД — пропускаем')
        return

    print(f'  Строк статистики ОРД: {len(rows)}')
    if dry_run:
        for r in rows:
            print(f'    {r["erid"]} | {r["platform"]} | показов: {r["shows"]}')
        return

    # Шаблон XLSX
    xlsx_path = build_ord_xlsx(rows, period_str)
    ord_dir = f'{accounting_folder}/ОРД/{period_str}'
    ensure_yadisk_folder(token, ord_dir)
    upload_file_to_yadisk(token, xlsx_path, f'{ord_dir}/Шаблон_ОРД_{period_str}.xlsx')
    Path(xlsx_path).unlink(missing_ok=True)

    # Инструкция (создаём один раз, перезаписываем)
    instr_docx = build_ord_instruction()
    instr_pdf = docx_to_pdf(instr_docx)
    instr_dir = f'{accounting_folder}/ОРД'
    ensure_yadisk_folder(token, instr_dir)
    upload_file_to_yadisk(token, instr_docx, f'{instr_dir}/Инструкция_ОРД.docx')
    upload_file_to_yadisk(token, instr_pdf, f'{instr_dir}/Инструкция_ОРД.pdf')
    Path(instr_docx).unlink(missing_ok=True)
    Path(instr_pdf).unlink(missing_ok=True)

    print(f'  Готово: шаблон и инструкция загружены в {instr_dir}')


# ---------------------------------------------------------------------------
# VK ORD API (опционально)
# ---------------------------------------------------------------------------

def push_to_vk_ord(rows, token_ord):
    """
    Отправляет статистику в ВК ОРД API.
    Вызывается только с явным флагом --push-to-ord после подтверждения пользователя.
    """
    base_url = 'https://api.ord.vk.com/v1'
    headers = {'Authorization': f'Bearer {token_ord}', 'Content-Type': 'application/json'}

    print('\n=== Отправка в ВК ОРД API ===')
    print('Будут отправлены данные:')
    for r in rows:
        print(f'  ERID: {r["erid"]} | {r["platform"]} | показов: {r["shows"]}')

    confirm = input('\nПодтвердите отправку (y/n): ').strip().lower()
    if confirm != 'y':
        print('Отправка отменена.')
        return

    for row in rows:
        payload = {
            'erid': row['erid'],
            'platform': row['platform'],
            'impressions': row['shows'],
            'paid_impressions': row['paid_shows'],
            'date_start': row['date_from'],
            'date_end': row['date_to'],
        }
        if row.get('total_sum'):
            payload['amount'] = row['total_sum']
        if row.get('event_type'):
            payload['event_type'] = row['event_type']

        r = requests.post(f'{base_url}/statistics', json=payload, headers=headers, timeout=30)
        if r.status_code in (200, 201):
            print(f'  ✓ {row["erid"]} — принято')
        else:
            print(f'  ✗ {row["erid"]} — ошибка {r.status_code}: {r.text[:200]}')


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--month', required=True)
    parser.add_argument('--stats-path', required=True)
    parser.add_argument('--dry-run', action='store_true')
    parser.add_argument('--push-to-ord', action='store_true')
    args = parser.parse_args()

    year, month = map(int, args.month.split('-'))
    token = os.environ['YANDEX_DISK_TOKEN']
    folder = os.environ.get('YANDEX_DISK_ACCOUNTING_FOLDER', 'Бухгалтерия')

    run(year, month, token, folder, args.stats_path, dry_run=args.dry_run)

    if args.push_to_ord:
        ord_token = os.environ.get('VK_ORD_API_TOKEN')
        if not ord_token:
            print('VK_ORD_API_TOKEN не задан — пропускаем API отправку')
        else:
            import openpyxl
            excel_bytes = download_excel_from_yadisk(token, args.stats_path)
            wb = openpyxl.load_workbook(io.BytesIO(excel_bytes))
            rows = parse_ord_sheet(wb)
            push_to_vk_ord(rows, ord_token)
